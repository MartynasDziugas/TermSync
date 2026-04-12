from pathlib import Path
from sqlalchemy.orm import Session
from docx import Document
from src.database.models import TemplateSegment, TermChangeType
from src.database.queries import get_changes_between_versions


def export_to_tmx(
    session: Session,
    old_template_id: int,
    new_template_id: int,
    output_path: str | Path,
) -> Path:
    try:
        old_segments = (
            session.query(TemplateSegment)
            .filter_by(template_id=old_template_id)
            .order_by(TemplateSegment.segment_index)
            .all()
        )
        new_segments = (
            session.query(TemplateSegment)
            .filter_by(template_id=new_template_id)
            .order_by(TemplateSegment.segment_index)
            .all()
        )

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        lines = ['<?xml version="1.0" encoding="UTF-8"?>',
                 '<tmx version="1.4"><body>']

        for i, new_seg in enumerate(new_segments):
            if i < len(old_segments):
                old_text = old_segments[i].source_text
                new_text = new_seg.source_text
                if old_text != new_text:
                    lines.append(f'  <tu>')
                    lines.append(f'    <tuv xml:lang="en-old"><seg>{old_text}</seg></tuv>')
                    lines.append(f'    <tuv xml:lang="en-new"><seg>{new_text}</seg></tuv>')
                    lines.append(f'  </tu>')

        lines.append('</body></tmx>')
        output_path.write_text("\n".join(lines), encoding="utf-8")
        return output_path
    except Exception as e:
        raise RuntimeError(f"Klaida eksportuojant į TMX: {e}")


def export_to_docx(
    session: Session,
    old_template_id: int,
    new_template_id: int,
    output_path: str | Path,
) -> Path:
    try:
        old_segments = (
            session.query(TemplateSegment)
            .filter_by(template_id=old_template_id)
            .order_by(TemplateSegment.segment_index)
            .all()
        )
        new_segments = (
            session.query(TemplateSegment)
            .filter_by(template_id=new_template_id)
            .order_by(TemplateSegment.segment_index)
            .all()
        )

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.add_heading("QRD šablonų pakeitimų ataskaita", level=1)
        doc.add_heading("Pasikeitę segmentai", level=2)

        changed = 0
        for i, new_seg in enumerate(new_segments):
            if i < len(old_segments):
                old_text = old_segments[i].source_text
                new_text = new_seg.source_text
                if old_text != new_text:
                    doc.add_paragraph(f"SENAS: {old_text}")
                    doc.add_paragraph(f"NAUJAS: {new_text}")
                    doc.add_paragraph("---")
                    changed += 1

        doc.add_paragraph(f"Iš viso pasikeitimų: {changed}")
        doc.save(output_path)
        return output_path
    except Exception as e:
        raise RuntimeError(f"Klaida eksportuojant į DOCX: {e}")