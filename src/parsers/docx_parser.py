from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def _is_likely_table_header_row(texts: list[str]) -> bool:
    """Ar pirmoji eilutė panaši į CAT / MemoQ antraštę (ne vertimo segmentas)."""
    if not texts:
        return False
    c0 = texts[0].strip().lower()
    if c0 == "id":
        return True
    lowered = [t.strip().lower() for t in texts if t.strip()]
    has_src = any("source" in t and "(en)" in t for t in lowered)
    has_tgt = any("target" in t and "(lt)" in t for t in lowered)
    if has_src and has_tgt:
        return True
    if len(texts) >= 2 and all(len(t.strip()) <= 24 for t in texts[:2]):
        a, b = texts[0].strip().lower(), texts[1].strip().lower()
        if a in ("en", "english", "source", "source (en)") and b in (
            "lt",
            "lithuanian",
            "target",
            "target (lt)",
        ):
            return True
    return False


def _detect_source_target_cols(texts: list[str]) -> tuple[int, int] | None:
    """I antraštės eilutės randa Source (en) ir Target (lt) stulpelių indeksus (0-based)."""
    src_i: int | None = None
    tgt_i: int | None = None
    for i, raw in enumerate(texts):
        t = raw.strip().lower()
        if src_i is None and "source" in t and "(en)" in t:
            src_i = i
        if tgt_i is None and "target" in t and "(lt)" in t:
            tgt_i = i
    if src_i is None or tgt_i is None or src_i == tgt_i:
        return None
    return src_i, tgt_i


def _cell_text(cell) -> str:
    """Lentelės ląstelės tekstas (visos pastraipos ląstelėje)."""
    parts: list[str] = []
    for para in cell.paragraphs:
        text = _paragraph_xml_text(para).strip()
        if not text:
            text = (para.text or "").strip()
        if text:
            parts.append(text)
    return " ".join(parts).strip()


def _paragraph_xml_text(paragraph) -> str:
    """
    Surenka visą pastraipos tekstą iš w:t elementų (įskaitant hipernuorodų ir laukų vidų),
    kur para.text kartais būna trumpesnis.
    """
    parts: list[str] = []
    for node in paragraph._p.iter(qn("w:t")):
        if node.text:
            parts.append(node.text)
    return "".join(parts)


class DocxParser:
    def __init__(self, file_path: str | Path) -> None:
        self.file_path = Path(file_path)

    def validate(self) -> bool:
        try:
            return self.file_path.exists() and self.file_path.suffix == ".docx"
        except Exception:
            return False

    def extract_segments(self) -> list[str]:
        try:
            doc = Document(self.file_path)
            segments: list[str] = []
            for para in doc.paragraphs:
                text = _paragraph_xml_text(para).strip()
                if not text:
                    text = (para.text or "").strip()
                if text:
                    segments.append(text)
            return segments
        except Exception as e:
            raise RuntimeError(f"Klaida skaitant .docx faila: {e}")

    def extract_bilingual_table_pairs(
        self,
        *,
        source_col: int = 0,
        target_col: int = 1,
        skip_header: bool = True,
        smart_column_layout: bool = True,
    ) -> list[tuple[str, str]]:
        """
        EN/LT poros iš Word lentelių. `source_col` / `target_col` — 0-based stulpelių indeksai.

        Jei `smart_column_layout` (numatytai True): aptinkama MemoQ/CAT antraštės eilutė
        (pvz. stulpeliai „Source (en)“, „Target (lt)“), stulpelių indeksai perkeliami į
        vėlesnes to paties failo lenteles su daug stulpelių. Antraštės eilutė praleidžiama
        tik jei ji tikrai panaši į antraštę (`skip_header`).

        Jei `smart_column_layout` yra False: visada naudojami tik `source_col` ir
        `target_col`; antraštės eilutė vis tiek praleidžiama tik pagal tą patį
        „header“ euristikos testą (kai `skip_header`).
        """
        if source_col < 0 or target_col < 0:
            raise ValueError("Stulpelių indeksai turi būti >= 0")
        if source_col == target_col:
            raise ValueError("Šaltinio ir tikslo stulpeliai turi skirtis")
        try:
            doc = Document(self.file_path)
            pairs: list[tuple[str, str]] = []
            last_src_tgt: tuple[int, int] | None = None

            for table in doc.tables:
                if not table.rows:
                    continue
                row0 = table.rows[0]
                row0_texts = [_cell_text(c) for c in row0.cells]
                ncols = len(row0.cells)

                header_like = bool(skip_header and _is_likely_table_header_row(row0_texts))
                if header_like and smart_column_layout:
                    detected = _detect_source_target_cols(row0_texts)
                    if detected is not None:
                        last_src_tgt = detected

                if smart_column_layout and last_src_tgt is not None:
                    src_c, tgt_c = last_src_tgt
                    if ncols > max(src_c, tgt_c):
                        eff_src, eff_tgt = src_c, tgt_c
                    else:
                        eff_src, eff_tgt = source_col, target_col
                else:
                    eff_src, eff_tgt = source_col, target_col

                max_col = max(eff_src, eff_tgt)
                start_ri = 1 if header_like else 0

                for ri in range(start_ri, len(table.rows)):
                    cells = table.rows[ri].cells
                    if len(cells) <= max_col:
                        continue
                    src = _cell_text(cells[eff_src])
                    tgt = _cell_text(cells[eff_tgt])
                    if not src or not tgt:
                        continue
                    pairs.append((src, tgt))

            return pairs
        except Exception as e:
            raise RuntimeError(f"Klaida skaitant lenteles .docx faile: {e}")
